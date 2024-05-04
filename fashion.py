from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
def truncate_text(text, num_words):
    words = text.split()
    truncated_words = words[:num_words]
    truncated_text = ' '.join(truncated_words)
    return truncated_text



driver = webdriver.Chrome()
driver.get("https://www.elle.com/fashion/")
elements = driver.find_elements(By.CSS_SELECTOR,"a[data-vars-cta='Personal Style']")
hrefs = []
for i in elements:
    hrefs.append(i.get_attribute('href'))
print(hrefs)
save_number = 0
for index in range(len(hrefs)):
    doc = Document()
    print(hrefs[index])
    driver.get(hrefs[index])
    try:
        WebDriverWait(driver,30).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR,"div[class^='article-body']")))
    except:
        WebDriverWait(driver,30).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR,"div[data-journey-body='listicle']")))
    try:
        article = driver.find_element(By.CSS_SELECTOR,"div[class^='article-body']").text
    except:
         article = driver.find_element(By.CSS_SELECTOR,"div[data-journey-body='listicle']").text
    truncated_text = truncate_text(article, 500)

    doc.add_paragraph(truncated_text)
    doc.save("scrapedDocuments/"+"Document-"+str(save_number+1)+"_FashionAndBeauty.docx")
    save_number = save_number + 1


hrefs = ["https://www.elle.com/beauty/makeup-skin-care/a60595170/climate-change-skin-care/","https://www.elle.com/beauty/makeup-skin-care/a60594081/pre-shower-makeup/","https://www.elle.com/beauty/makeup-skin-care/a60560119/emily-ratajkowski-lipstick/"]

for index in range(len(hrefs)):
    doc = Document()
    print(index)
    driver.get(hrefs[index])
    WebDriverWait(driver,30).until(EC.presence_of_all_elements_located((By.CSS_SELECTOR,"div[class^='article-body']")))
    article = driver.find_element(By.CSS_SELECTOR,"div[class^='article-body']").text
    truncated_text = truncate_text(article, 600)

    doc.add_paragraph(truncated_text)
    doc.save("scrapedDocuments/"+"Document-"+str(save_number+1)+"_FashionAndBeauty.docx")
    save_number = save_number + 1
